# -*- coding: utf-8 -*-
"""模板相关接口：列表、创建、详情、更新、删除、发布、二维码、解析、名单关联。"""
import os
import json
import base64
from io import BytesIO
from datetime import datetime, timedelta

from flask import Blueprint, request, jsonify, send_file
from docx import Document

from ..config import TPL_DIR, get_fill_base_url, _is_private_host
from ..db import get_db, gen_id, now_str
from ..utils import (infer_type, docx_to_type, extract_placeholders,
                     extract_placeholders_xlsx, safe_filename, tpl_to_dict)

bp = Blueprint('templates', __name__)


@bp.route('/api/templates', methods=['GET'])
def api_list_templates():
    db = get_db()
    rows = db.execute('SELECT * FROM templates WHERE deleted = 0 ORDER BY updated_at DESC').fetchall()
    return jsonify({'code': 0, 'data': [tpl_to_dict(r) for r in rows]})


@bp.route('/api/templates', methods=['POST'])
def api_create_template():
    """通过已解析的字段创建模板（前端已完成 docx 解析）。"""
    data = request.get_json() or {}
    name = (data.get('name') or '').strip()
    if not name:
        return jsonify({'code': 1, 'message': '模板名称不能为空'}), 400
    fields = data.get('fields') or []
    file_name = data.get('file_name') or '未上传文件.docx'
    file_b64 = data.get('file_base64')

    tpl_id = gen_id()
    file_path = os.path.join(TPL_DIR, f'{tpl_id}_{safe_filename(file_name)}')

    if file_b64:
        try:
            with open(file_path, 'wb') as f:
                f.write(base64.b64decode(file_b64))
        except Exception as e:
            return jsonify({'code': 1, 'message': f'文件保存失败: {e}'}), 500
    else:
        try:
            doc = Document()
            doc.add_paragraph(f'模板: {name}')
            for f in fields:
                doc.add_paragraph(f'{{{{{f.get("name", "")}}}}}')
            doc.save(file_path)
        except Exception as e:
            return jsonify({'code': 1, 'message': f'占位文件创建失败: {e}'}), 500

    db = get_db()
    db.execute('''INSERT INTO templates (id, name, category, file_name, file_path, fields_json, status, created_at, updated_at)
                  VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)''',
               (tpl_id, name, data.get('category', '未分类'), file_name, file_path,
                json.dumps(fields, ensure_ascii=False), data.get('status', 'draft'),
                now_str(), now_str()))
    db.commit()
    row = db.execute('SELECT * FROM templates WHERE id = ?', (tpl_id,)).fetchone()
    return jsonify({'code': 0, 'data': tpl_to_dict(row)})


@bp.route('/api/templates/<tid>', methods=['GET'])
def api_get_template(tid):
    db = get_db()
    row = db.execute('SELECT * FROM templates WHERE id = ?', (tid,)).fetchone()
    if not row:
        return jsonify({'code': 1, 'message': '模板不存在'}), 404
    return jsonify({'code': 0, 'data': tpl_to_dict(row)})


@bp.route('/api/templates/<tid>', methods=['PUT'])
def api_update_template(tid):
    data = request.get_json() or {}
    db = get_db()
    row = db.execute('SELECT * FROM templates WHERE id = ?', (tid,)).fetchone()
    if not row:
        return jsonify({'code': 1, 'message': '模板不存在'}), 404
    fields = data.get('fields')
    name = data.get('name')
    category = data.get('category')
    status = data.get('status')
    published_at = row['published_at']
    if status == 'published' and not published_at:
        published_at = now_str()
    # 支持更新 Word 模板文件
    file_b64 = data.get('file_base64')
    file_name = data.get('file_name')
    if file_b64 and file_name:
        new_path = os.path.join(TPL_DIR, f'{tid}_{safe_filename(file_name)}')
        try:
            with open(new_path, 'wb') as f:
                f.write(base64.b64decode(file_b64))
            db.execute('UPDATE templates SET file_name = ?, file_path = ? WHERE id = ?',
                       (file_name, new_path, tid))
        except Exception:
            pass
    db.execute('''UPDATE templates SET
                    name = COALESCE(?, name),
                    category = COALESCE(?, category),
                    fields_json = COALESCE(?, fields_json),
                    status = COALESCE(?, status),
                    published_at = ?,
                    updated_at = ?
                  WHERE id = ?''',
               (name, category, json.dumps(fields, ensure_ascii=False) if fields else None,
                status, published_at, now_str(), tid))
    db.commit()
    row = db.execute('SELECT * FROM templates WHERE id = ?', (tid,)).fetchone()
    return jsonify({'code': 0, 'data': tpl_to_dict(row)})


@bp.route('/api/templates/<tid>', methods=['DELETE'])
def api_delete_template(tid):
    db = get_db()
    row = db.execute('SELECT * FROM templates WHERE id = ?', (tid,)).fetchone()
    if not row:
        return jsonify({'code': 1, 'message': '模板不存在'}), 404
    rid = gen_id()
    expire = (datetime.now() + timedelta(days=30)).strftime('%Y-%m-%d %H:%M:%S')
    db.execute('''INSERT INTO recycle (id, item_type, item_id, item_name, deleted_at, expire_at)
                  VALUES (?, ?, ?, ?, ?, ?)''',
               (rid, 'template', tid, row['name'], now_str(), expire))
    db.execute('UPDATE templates SET deleted = 1 WHERE id = ?', (tid,))
    db.commit()
    return jsonify({'code': 0, 'data': True})


@bp.route('/api/templates/<tid>/publish', methods=['POST'])
def api_publish_template(tid):
    """切换发布状态：已发布 -> 草稿（下架），草稿 -> 已发布（上架）。

    也可通过请求体 {"status": "published"|"draft"} 显式指定目标状态。
    """
    db = get_db()
    row = db.execute('SELECT * FROM templates WHERE id = ?', (tid,)).fetchone()
    if not row:
        return jsonify({'code': 1, 'message': '模板不存在'}), 404

    payload = request.get_json(silent=True) or {}
    target = payload.get('status')
    if target not in ('published', 'draft'):
        # 未显式指定则按当前状态取反
        target = 'draft' if row['status'] == 'published' else 'published'

    if target == 'published':
        db.execute('UPDATE templates SET status = ?, published_at = ?, updated_at = ? WHERE id = ?',
                   (target, now_str(), now_str(), tid))
    else:
        db.execute('UPDATE templates SET status = ?, updated_at = ? WHERE id = ?',
                   (target, now_str(), tid))
    db.commit()
    row = db.execute('SELECT * FROM templates WHERE id = ?', (tid,)).fetchone()
    return jsonify({'code': 0, 'data': tpl_to_dict(row)})


@bp.route('/api/templates/<tid>/qrcode', methods=['GET'])
def api_template_qrcode(tid):
    """生成模板填写页二维码（微信扫码）。

    支持 ?base=… 覆盖 base URL（前端 `window.location.origin`）。
    不传时回退到 get_fill_base_url()（自动适配站点设置/请求 host）。
    """
    db = get_db()
    tpl = db.execute('SELECT * FROM templates WHERE id = ?', (tid,)).fetchone()
    if not tpl:
        return jsonify({'code': 1, 'message': '模板不存在'}), 404

    base_url = (request.args.get('base') or '').strip().rstrip('/')
    if not base_url:
        base_url = get_fill_base_url()
    # 安全兜底：拒绝「不可达」的内网/容器私有地址（如 10.255.254.20 这类隧道/容器虚拟 IP），
    # 但允许本地回环(127.0.0.1 / localhost)以便本机开发调试。
    if _is_private_host(base_url):
        host_part = base_url.split('://', 1)[-1].split('/')[0].split(':')[0].lower()
        if host_part not in ('127.0.0.1', 'localhost', '::1'):
            base_url = ''
    if not base_url:
        return jsonify({'code': 1, 'message': '无法生成公网访问链接：请在「系统设置」填写站点公网地址，或在部署平台设置环境变量 BASE_URL'}), 400
    fill_url = f'{base_url}/fill/{tid}'

    import qrcode
    from qrcode.image.pure import PyPNGImage
    qr = qrcode.QRCode(version=1, box_size=8, border=2)
    qr.add_data(fill_url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="#1a2634", back_color="white", image_factory=PyPNGImage)
    buf = BytesIO()
    img.save(buf)
    b64 = base64.b64encode(buf.getvalue()).decode('ascii')

    link = db.execute('SELECT * FROM template_roster WHERE template_id = ?', (tid,)).fetchone()
    roster_info = None
    if link:
        roster = db.execute('SELECT * FROM rosters WHERE id = ?', (link['roster_id'],)).fetchone()
        if roster:
            roster_info = {
                'roster_name': roster['name'],
                'roster_total': roster['total'],
                'id_field': link['id_field'],
                'name_field': link['name_field'],
            }

    return jsonify({'code': 0, 'data': {
        'image': f'data:image/png;base64,{b64}',
        'url': fill_url,
        'template_name': tpl['name'],
        'template_status': tpl['status'],
        'roster': roster_info,
    }})


@bp.route('/api/templates/parse', methods=['POST'])
def api_parse_template():
    """服务器端解析（前端已实现解析，此接口作为备用）。"""
    if 'file' not in request.files:
        return jsonify({'code': 1, 'message': '未提供文件'}), 400
    f = request.files['file']
    file_bytes = f.read()
    if f.filename.lower().endswith('.xlsx'):
        names = extract_placeholders_xlsx(file_bytes)
    else:
        names = extract_placeholders(file_bytes)
    fields = []
    for n in names:
        info = infer_type(n)
        fields.append({
            'name': n,
            'type': docx_to_type(info.get('type', 'text')),
            'raw_type': info.get('type', 'text'),
            'options': ','.join(info.get('options', [])),
            'pattern': info.get('pattern', ''),
            'placeholder': info.get('placeholder', ''),
            'hint': info.get('hint', ''),
            'required': True,
            'unique': False,
        })
    return jsonify({'code': 0, 'data': {
        'fileName': f.filename,
        'verified': len(fields) > 0,
        'fields': fields,
    }})


@bp.route('/api/templates/<tid>/link-roster', methods=['POST'])
def api_link_roster(tid):
    """将名单关联到模板。"""
    data = request.get_json() or {}
    rid = data.get('roster_id')
    id_field = data.get('id_field', '学号')
    name_field = data.get('name_field', '姓名')
    if not rid:
        return jsonify({'code': 1, 'message': '缺少 roster_id'}), 400
    db = get_db()
    tpl = db.execute('SELECT * FROM templates WHERE id = ?', (tid,)).fetchone()
    if not tpl:
        return jsonify({'code': 1, 'message': '模板不存在'}), 404
    roster = db.execute('SELECT * FROM rosters WHERE id = ?', (rid,)).fetchone()
    if not roster:
        return jsonify({'code': 1, 'message': '名单不存在'}), 404
    db.execute('''INSERT OR REPLACE INTO template_roster (template_id, roster_id, id_field, name_field, linked_at)
                  VALUES (?, ?, ?, ?, ?)''', (tid, rid, id_field, name_field, now_str()))
    db.commit()
    return jsonify({'code': 0, 'data': {
        'template_id': tid, 'roster_id': rid,
        'id_field': id_field, 'name_field': name_field,
        'roster_name': roster['name'], 'roster_total': roster['total'],
    }})


@bp.route('/api/templates/<tid>/roster', methods=['GET'])
def api_get_template_roster(tid):
    """获取模板关联的名单信息。"""
    db = get_db()
    link = db.execute('SELECT * FROM template_roster WHERE template_id = ?', (tid,)).fetchone()
    if not link:
        return jsonify({'code': 0, 'data': None})
    roster = db.execute('SELECT * FROM rosters WHERE id = ?', (link['roster_id'],)).fetchone()
    if not roster:
        return jsonify({'code': 0, 'data': None})
    return jsonify({'code': 0, 'data': {
        'roster_id': roster['id'],
        'roster_name': roster['name'],
        'roster_total': roster['total'],
        'id_field': link['id_field'],
        'name_field': link['name_field'],
    }})


@bp.route('/api/templates/<tid>/roster-progress', methods=['GET'])
def api_roster_progress(tid):
    """获取模板的名单提交进度：谁已提交、谁未提交。"""
    db = get_db()
    link = db.execute('SELECT * FROM template_roster WHERE template_id = ?', (tid,)).fetchone()
    if not link:
        return jsonify({'code': 0, 'data': None})
    roster = db.execute('SELECT * FROM rosters WHERE id = ?', (link['roster_id'],)).fetchone()
    if not roster:
        return jsonify({'code': 0, 'data': None})

    id_field = link['id_field']
    name_field = link['name_field']
    try:
        rows = json.loads(roster['rows_json'])
    except Exception:
        rows = []

    subs = db.execute('SELECT * FROM submissions WHERE template_id = ?', (tid,)).fetchall()
    # 按名额行 row_id 定位提交状态（支持同一人多个名额各自独立）
    sub_by_row = {}
    for s in subs:
        rid = s['roster_row_id']
        if rid:
            sub_by_row[rid] = {
                'submitter': s['submitter'] or '',
                'submitted_at': s['submitted_at'],
                'submission_id': s['id'],
            }

    members, submitted_count = [], 0
    for r in rows:
        rid = str(r.get('row_id', '')).strip()
        sid = str(r.get(id_field, '')).strip()
        sname = str(r.get(name_field, '')).strip()
        if not sid and not sname:
            continue
        info = sub_by_row.get(rid, {})
        is_submitted = bool(info)
        if is_submitted:
            submitted_count += 1
        members.append({
            'id': sid,
            'name': sname,
            'row_id': rid,
            'submitted': is_submitted,
            'info': info,
            'extra': {k: v for k, v in r.items() if k not in (id_field, name_field, 'row_id')},
        })

    total = len(members)
    return jsonify({'code': 0, 'data': {
        'roster_name': roster['name'],
        'id_field': id_field,
        'name_field': name_field,
        'total': total,
        'submitted': submitted_count,
        'unsubmitted': total - submitted_count,
        'progress': round(submitted_count / total * 100, 1) if total > 0 else 0,
        'members': members,
    }})
